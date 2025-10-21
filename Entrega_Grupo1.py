"""
Script para gerar documento Word da entrega do grupo
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ====================
# TÍTULO
# ====================
titulo = doc.add_heading('Análise de Gastos Parlamentares', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
subtitulo = doc.add_paragraph('Câmara dos Deputados - Cota Parlamentar')
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitulo.runs[0].font.size = Pt(14)
subtitulo.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()  # Espaço

# ====================
# NÚMERO DO GRUPO
# ====================
doc.add_heading('1. Número do Grupo', 1)
p = doc.add_paragraph('Grupo: ')
p.add_run('1').bold = True
p.runs[1].font.size = Pt(14)

# ====================
# INTEGRANTES
# ====================
doc.add_heading('2. Integrantes do Grupo', 1)

integrantes = [
    ('Leticia', '21352'),
    ('Gabriel', '24734'),
    ('Thales', '24740'),
    ('Maria Fernanda', '24767')
]

for nome, matricula in integrantes:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{nome} - ').font.size = Pt(12)
    p.add_run(f'Matrícula {matricula}').font.size = Pt(12)

# ====================
# BIBLIOTECAS
# ====================
doc.add_heading('3. Bibliotecas Python Previstas para Uso', 1)

bibliotecas = [
    {
        'nome': 'Pandas',
        'versao': '2.3.3',
        'justificativa': 'Manipulação e análise de dados estruturados. Essencial para leitura de CSV, limpeza de dados, agregações estatísticas e cruzamento de informações entre diferentes fontes de dados.'
    },
    {
        'nome': 'Requests',
        'versao': '2.32.5',
        'justificativa': 'Comunicação com APIs REST. Utilizado para fazer requisições HTTP à API da Câmara dos Deputados e obter dados cadastrais dos deputados (partido, estado, etc.).'
    },
    {
        'nome': 'Matplotlib',
        'versao': '3.10.7',
        'justificativa': 'Biblioteca base para visualização de dados em Python. Permite criar gráficos personalizados em alta resolução (300 DPI) com controle total sobre elementos visuais.'
    },
    {
        'nome': 'Seaborn',
        'versao': '0.13.2',
        'justificativa': 'Extensão do Matplotlib com templates profissionais. Facilita criação de gráficos estatísticos complexos com estética moderna e cores otimizadas.'
    },
    {
        'nome': 'NumPy',
        'versao': '2.3.4',
        'justificativa': 'Operações numéricas e arrays multidimensionais. Dependência do Pandas, utilizado para cálculos estatísticos eficientes e operações matemáticas vetorizadas.'
    },
    {
        'nome': 'Unidecode',
        'versao': '1.4.0',
        'justificativa': 'Normalização de texto removendo acentos e caracteres especiais. Fundamental para padronizar nomes de deputados e garantir alta taxa de match (>95%) no cruzamento de dados.'
    },
    {
        'nome': 'python-pptx',
        'versao': '1.0.2',
        'justificativa': 'Geração automática de apresentações PowerPoint. Cria slides profissionais com gráficos, tabelas e insights, automatizando a preparação do material de apresentação.'
    },
    {
        'nome': 'python-docx',
        'versao': '1.2.0',
        'justificativa': 'Criação e manipulação de documentos Word (.docx). Utilizado para gerar documentação e relatórios formatados programaticamente.'
    }
]

for i, bib in enumerate(bibliotecas, 1):
    # Nome da biblioteca
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {bib['nome']}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Versão
    p.add_run(f" (v{bib['versao']})")
    p.runs[-1].font.size = Pt(11)
    p.runs[-1].italic = True
    
    # Justificativa
    p_just = doc.add_paragraph(bib['justificativa'])
    p_just.paragraph_format.left_indent = Inches(0.3)
    p_just.runs[0].font.size = Pt(11)

# ====================
# ANÁLISES E VISUALIZAÇÕES
# ====================
doc.add_page_break()
doc.add_heading('4. Análises e Visualizações Previstas', 1)

analises = [
    {
        'titulo': 'Gastos por Partido Político',
        'tipo': 'Gráfico de Barras Horizontal',
        'objetivo': 'Comparar o volume total de gastos entre diferentes partidos políticos. Permite identificar quais partidos têm maior volume de despesas parlamentares e analisar se há correlação com o tamanho da bancada.',
        'metricas': ['Valor total gasto por partido', 'Quantidade de deputados por partido', 'Gasto médio por deputado']
    },
    {
        'titulo': 'Gastos por Estado (UF)',
        'tipo': 'Gráfico de Barras Horizontal',
        'objetivo': 'Visualizar a distribuição geográfica dos gastos parlamentares. Identificar estados com maiores volumes de despesas e verificar se há relação com fatores como população, número de deputados ou distância da capital federal.',
        'metricas': ['Valor total gasto por estado', 'Número de deputados por estado', 'Gasto médio por deputado estadual']
    },
    {
        'titulo': 'Principais Tipos de Despesa',
        'tipo': 'Gráfico de Barras Horizontal (Top 15)',
        'objetivo': 'Identificar as categorias de despesas mais representativas no orçamento parlamentar. Compreender onde se concentram os principais gastos da cota parlamentar e quais tipos de despesas são mais frequentes.',
        'metricas': ['Valor total por tipo de despesa', 'Percentual de cada tipo no total', 'Frequência de uso de cada categoria']
    },
    {
        'titulo': 'Top 20 Deputados com Maiores Gastos',
        'tipo': 'Gráfico de Barras Horizontal',
        'objetivo': 'Ranquear os deputados com maiores volumes de gastos individuais. Permite análise de outliers e identificação de padrões de uso da cota parlamentar. Inclui informações de partido e estado para análise contextual.',
        'metricas': ['Valor total gasto por deputado', 'Partido e estado do deputado', 'Ranking comparativo']
    },
    {
        'titulo': 'Dashboard - Resumo Geral',
        'tipo': 'Painel com 4 Visualizações',
        'objetivo': 'Fornecer uma visão consolidada das principais análises em um único gráfico. Facilita a compreensão global dos dados e permite comparações rápidas entre diferentes dimensões de análise (partido, estado, despesas, deputados).',
        'metricas': ['Compilação dos 4 gráficos principais', 'Layout otimizado para apresentação', 'Visão panorâmica dos dados']
    }
]

for i, analise in enumerate(analises, 1):
    # Título da análise
    heading = doc.add_heading(f"{i}. {analise['titulo']}", 2)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    # Tipo de visualização
    p = doc.add_paragraph()
    p.add_run('Tipo de Visualização: ').bold = True
    p.add_run(analise['tipo'])
    p.runs[0].font.size = Pt(11)
    p.runs[1].font.size = Pt(11)
    
    # Objetivo
    p = doc.add_paragraph()
    p.add_run('Objetivo: ').bold = True
    p.add_run(analise['objetivo'])
    p.runs[0].font.size = Pt(11)
    p.runs[1].font.size = Pt(11)
    
    # Métricas
    p = doc.add_paragraph()
    p.add_run('Métricas Principais:').bold = True
    p.runs[0].font.size = Pt(11)
    
    for metrica in analise['metricas']:
        p_met = doc.add_paragraph(metrica, style='List Bullet')
        p_met.paragraph_format.left_indent = Inches(0.5)
        p_met.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Espaço entre análises

# ====================
# METODOLOGIA
# ====================
doc.add_page_break()
doc.add_heading('5. Metodologia de Análise', 1)

metodologia_texto = """
O projeto utiliza uma abordagem sistemática de análise de dados governamentais em 6 etapas:

1. Carregamento e Limpeza de Dados
   - Leitura de arquivo CSV com dados de despesas parlamentares
   - Remoção de registros com valores nulos ou inválidos (≤ 0)
   - Padronização de nomes (uppercase, remoção de acentos)
   - Taxa de aproveitamento de dados: ~64% dos registros originais

2. Integração com API da Câmara dos Deputados
   - Requisições à API REST oficial (https://dadosabertos.camara.leg.br/api/v2/deputados)
   - Coleta de dados cadastrais: nome completo, partido, estado (UF)
   - Tratamento de erros e timeouts de rede

3. Cruzamento de Dados
   - Match por nome normalizado (uppercase + sem acentos)
   - Taxa de identificação: >95% dos deputados
   - Enriquecimento de dados com informações partidárias e estaduais

4. Análise Estatística
   - Agregações por múltiplas dimensões (partido, estado, tipo de despesa)
   - Cálculos de totais, médias e rankings
   - Identificação de outliers e padrões

5. Geração de Visualizações
   - Criação de 5 gráficos profissionais em PNG (300 DPI)
   - Paleta de cores consistente e acessível
   - Layout otimizado para apresentações

6. Geração Automática de Apresentação
   - Criação de PowerPoint completo com 15 slides
   - Integração automática de todos os gráficos e insights
   - Inclusão de dados estatísticos e rankings

Toda execução é organizada em pasta timestampada (execucao_YYYYMMDD_HHMMSS/) contendo:
- 5 arquivos CSV com dados processados
- 5 gráficos PNG em alta resolução
- 1 apresentação PowerPoint completa
"""

p = doc.add_paragraph(metodologia_texto)
p.runs[0].font.size = Pt(11)

# ====================
# RESULTADOS ESPERADOS
# ====================
doc.add_heading('6. Resultados Esperados', 1)

resultados_texto = """
Ao final da análise, espera-se:

✓ Identificação Clara de Padrões de Gastos
  - Quais partidos têm maior volume de despesas
  - Quais estados concentram mais gastos
  - Quais tipos de despesas são mais comuns

✓ Rankings e Comparações
  - Top 20 deputados com maiores gastos individuais
  - Comparação entre bancadas partidárias
  - Distribuição geográfica de despesas

✓ Insights para Análise Política
  - Correlação entre tamanho de bancada e volume de gastos
  - Diferenças regionais nos padrões de despesas
  - Tipos de despesas prioritários por categoria

✓ Material de Apresentação Profissional
  - Visualizações de alta qualidade para relatórios
  - Dados estruturados para análises adicionais
  - Apresentação PowerPoint automatizada e completa

✓ Reprodutibilidade e Automação
  - Pipeline automatizado de ponta a ponta
  - Organização clara de resultados por execução
  - Código modular e reutilizável para novos datasets
"""

p = doc.add_paragraph(resultados_texto)
p.runs[0].font.size = Pt(11)

# ====================
# RODAPÉ
# ====================
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph('_______________________________________________')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

info = doc.add_paragraph('Ciência de Dados - Análise de Dados Governamentais')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(10)
info.runs[0].italic = True

data = doc.add_paragraph('Outubro de 2025')
data.alignment = WD_ALIGN_PARAGRAPH.CENTER
data.runs[0].font.size = Pt(10)
data.runs[0].italic = True

# Salvar documento
output_path = r'c:\Users\thale\Docs\Faculdade\Ciencia de dados\Analise de dados do governo api\Entrega_Grupo1.docx'
doc.save(output_path)

print(f"✓ Documento criado com sucesso!")
print(f"📄 Salvo em: {output_path}")
print(f"\n📋 Conteúdo do documento:")
print("   • Número do Grupo: 1")
print("   • Integrantes: 4 membros com matrículas")
print("   • Bibliotecas: 8 bibliotecas com justificativas detalhadas")
print("   • Análises: 5 visualizações previstas com objetivos e métricas")
print("   • Metodologia: Processo completo em 6 etapas")
print("   • Resultados Esperados: Insights e entregas finais")
print(f"\n📄 Total de páginas: ~{6} páginas")
